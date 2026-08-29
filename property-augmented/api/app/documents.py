from __future__ import annotations

import hashlib
import io
import json
import shutil
import uuid
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from fastapi import Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from openpyxl import load_workbook
from pydantic import BaseModel, Field
from pypdf import PdfReader
from sqlalchemy import DateTime, ForeignKey, Integer, String, Text
from sqlalchemy.orm import Mapped, Session, mapped_column

from . import full_stack as stack
from . import main as core
from .agents import AgentRunRequest, _run, _security_flags
from .bootstrap import app

MAX_BYTES = 30 * 1024 * 1024
MAX_EXTRACTED_CHARS = 300_000
MAX_ZIP_ENTRIES = 3000
MAX_ZIP_UNCOMPRESSED = 150 * 1024 * 1024
MAX_PDF_PAGES = 750
ALLOWED = {".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".txt", ".md", ".json"}
MIMES = {
    ".pdf":"application/pdf",
    ".docx":"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx":"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".xlsm":"application/vnd.ms-excel.sheet.macroEnabled.12",
    ".csv":"text/csv",
    ".txt":"text/plain",
    ".md":"text/markdown",
    ".json":"application/json",
}


def now() -> datetime:
    return datetime.now(timezone.utc)


class DocumentAsset(stack.Base):
    __tablename__ = "pda_documents"
    id: Mapped[int] = mapped_column(Integer, primary_key=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("pda_users.id"), index=True)
    project_id: Mapped[int | None] = mapped_column(ForeignKey("pda_projects.id"), nullable=True, index=True)
    original_name: Mapped[str] = mapped_column(String(500))
    storage_key: Mapped[str] = mapped_column(String(700), unique=True)
    sha256: Mapped[str] = mapped_column(String(64), index=True)
    size_bytes: Mapped[int] = mapped_column(Integer)
    suffix: Mapped[str] = mapped_column(String(20))
    mime_type: Mapped[str] = mapped_column(String(160), default="application/octet-stream")
    extraction_status: Mapped[str] = mapped_column(String(60), default="pending")
    extracted_text: Mapped[str] = mapped_column(Text, default="")
    extraction_note: Mapped[str] = mapped_column(Text, default="")
    security_flags_json: Mapped[str] = mapped_column(Text, default="[]")
    retention_until: Mapped[datetime] = mapped_column(DateTime(timezone=True))
    created_at: Mapped[datetime] = mapped_column(DateTime(timezone=True), default=now)


stack.Base.metadata.create_all(stack.engine)


class AnalyseDocument(BaseModel):
    goal: str = Field(default="Audit this document for established facts, assumptions, dates, revisions, contradictions, omissions, risks and verification actions.", min_length=3, max_length=20000)
    web_research: bool = False
    allowed_domains: list[str] = Field(default_factory=list)


def _owned_project(project_id: int | None, user: stack.User, s: Session) -> None:
    if project_id is None:
        return
    p=s.get(stack.Project,project_id)
    if not p or p.user_id!=user.id:
        raise HTTPException(404,"Project not found")


def _path_for(user_id:int,project_id:int|None,suffix:str)->Path:
    folder=core.STORAGE/f"user-{user_id}"/(f"project-{project_id}" if project_id is not None else "unassigned")
    folder.mkdir(parents=True,exist_ok=True)
    return folder/f"{uuid.uuid4().hex}{suffix}"


def _inspect_zip(raw:bytes)->None:
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as z:
            entries=z.infolist()
            if len(entries)>MAX_ZIP_ENTRIES:raise HTTPException(422,"Office file contains too many archive entries")
            total=sum(x.file_size for x in entries)
            if total>MAX_ZIP_UNCOMPRESSED:raise HTTPException(422,"Office file expands beyond the safe processing limit")
            if any(x.flag_bits & 0x1 for x in entries):raise HTTPException(422,"Encrypted Office containers are not supported")
            for x in entries:
                if x.file_size>80*1024*1024:raise HTTPException(422,"Office file contains an oversized embedded item")
                if x.compress_size and x.file_size/x.compress_size>250:raise HTTPException(422,"Office file compression ratio exceeds the safe processing limit")
    except zipfile.BadZipFile as exc:
        raise HTTPException(422,"Invalid Office document container") from exc


def _extract(raw:bytes,suffix:str,path:Path)->tuple[str,str]:
    try:
        if suffix==".pdf":
            reader=PdfReader(io.BytesIO(raw))
            if reader.is_encrypted:raise HTTPException(422,"Encrypted PDFs must be decrypted before upload")
            if len(reader.pages)>MAX_PDF_PAGES:raise HTTPException(422,f"PDF exceeds the {MAX_PDF_PAGES}-page processing limit")
            text="\n".join((p.extract_text() or "") for p in reader.pages)
            return text[:MAX_EXTRACTED_CHARS],f"Extracted text from {len(reader.pages)} PDF pages. Image-only pages may require separate OCR/manual review."
        if suffix==".docx":
            _inspect_zip(raw);doc=DocxDocument(io.BytesIO(raw));text="\n".join(p.text for p in doc.paragraphs)
            return text[:MAX_EXTRACTED_CHARS],"Extracted paragraph text. Embedded images/drawings are not treated as text evidence."
        if suffix in {".xlsx",".xlsm"}:
            _inspect_zip(raw);wb=load_workbook(io.BytesIO(raw),data_only=True,read_only=True,keep_vba=False);lines=[]
            for ws in wb.worksheets:
                lines.append(f"[SHEET: {ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    lines.append("\t".join("" if v is None else str(v) for v in row))
                    if sum(len(x) for x in lines)>MAX_EXTRACTED_CHARS:break
                if sum(len(x) for x in lines)>MAX_EXTRACTED_CHARS:break
            return "\n".join(lines)[:MAX_EXTRACTED_CHARS],"Workbook values extracted with formulas evaluated only where cached values exist. Macros are never executed."
        text=raw.decode("utf-8",errors="replace")
        if suffix==".json":
            try:json.loads(text)
            except json.JSONDecodeError as exc:raise HTTPException(422,"Invalid JSON document") from exc
        return text[:MAX_EXTRACTED_CHARS],"Text decoded as UTF-8 with replacement for invalid bytes."
    except HTTPException:raise
    except Exception as exc:
        return "",f"Extraction failed safely: {type(exc).__name__}: {exc}"


def _serialise(x:DocumentAsset,include_preview:bool=False)->dict[str,Any]:
    out={"id":x.id,"project_id":x.project_id,"original_name":x.original_name,"sha256":x.sha256,"size_bytes":x.size_bytes,"suffix":x.suffix,"mime_type":x.mime_type,"extraction_status":x.extraction_status,"extraction_note":x.extraction_note,"security_flags":json.loads(x.security_flags_json or "[]"),"retention_until":x.retention_until,"created_at":x.created_at}
    if include_preview:out["text_preview"]=x.extracted_text[:4000]
    return out


@app.post("/api/v1/documents/secure-upload")
async def secure_upload(file:UploadFile=File(...),project_id:int|None=Form(default=None),retention_days:int=Form(default=90),user:stack.User=Depends(stack.me),s:Session=Depends(stack.db)):
    _owned_project(project_id,user,s)
    name=(file.filename or "document").strip();suffix=Path(name).suffix.lower()
    if suffix not in ALLOWED:raise HTTPException(415,"Unsupported document type")
    retention_days=max(1,min(int(retention_days),3650))
    raw=await file.read(MAX_BYTES+1)
    if len(raw)>MAX_BYTES:raise HTTPException(413,"30MB document limit")
    if not raw:raise HTTPException(422,"Empty document")
    path=_path_for(user.id,project_id,suffix);path.write_bytes(raw)
    try:text,note=_extract(raw,suffix,path)
    except Exception:
        path.unlink(missing_ok=True);raise
    status="extracted" if text else "extraction_limited"
    flags=_security_flags({"document_text":text})
    asset=DocumentAsset(user_id=user.id,project_id=project_id,original_name=name[:500],storage_key=str(path.relative_to(core.STORAGE)),sha256=hashlib.sha256(raw).hexdigest(),size_bytes=len(raw),suffix=suffix,mime_type=MIMES.get(suffix,file.content_type or "application/octet-stream"),extraction_status=status,extracted_text=text,extraction_note=note,security_flags_json=json.dumps(flags),retention_until=now()+timedelta(days=retention_days))
    s.add(asset);s.commit();s.refresh(asset)
    return _serialise(asset,True)


@app.get("/api/v1/documents")
def list_documents(project_id:int|None=None,user:stack.User=Depends(stack.me),s:Session=Depends(stack.db)):
    q=s.query(DocumentAsset).filter(DocumentAsset.user_id==user.id)
    if project_id is not None:q=q.filter(DocumentAsset.project_id==project_id)
    rows=q.order_by(DocumentAsset.created_at.desc()).limit(500).all()
    return {"documents":[_serialise(x) for x in rows]}


@app.get("/api/v1/documents/{document_id}/content")
def document_content(document_id:int,user:stack.User=Depends(stack.me),s:Session=Depends(stack.db)):
    x=s.get(DocumentAsset,document_id)
    if not x or x.user_id!=user.id:raise HTTPException(404,"Document not found")
    return {**_serialise(x),"extracted_text":x.extracted_text}


@app.get("/api/v1/documents/{document_id}/download")
def document_download(document_id:int,user:stack.User=Depends(stack.me),s:Session=Depends(stack.db)):
    x=s.get(DocumentAsset,document_id)
    if not x or x.user_id!=user.id:raise HTTPException(404,"Document not found")
    path=(core.STORAGE/x.storage_key).resolve();root=core.STORAGE.resolve()
    if root not in path.parents or not path.exists():raise HTTPException(404,"Stored file not found")
    return FileResponse(path,filename=x.original_name,media_type=x.mime_type)


@app.post("/api/v1/documents/{document_id}/analyse")
async def analyse_document(document_id:int,r:AnalyseDocument,user:stack.User=Depends(stack.me),s:Session=Depends(stack.db)):
    x=s.get(DocumentAsset,document_id)
    if not x or x.user_id!=user.id:raise HTTPException(404,"Document not found")
    context={"document":{"id":x.id,"filename":x.original_name,"sha256":x.sha256,"size_bytes":x.size_bytes,"extraction_status":x.extraction_status,"extraction_note":x.extraction_note,"security_flags":json.loads(x.security_flags_json or "[]"),"extracted_text":x.extracted_text}}
    return await _run("document-auditor",AgentRunRequest(goal=r.goal,project_id=x.project_id,context=context,web_research=r.web_research,allowed_domains=r.allowed_domains),user,s)


@app.delete("/api/v1/documents/{document_id}")
def delete_document(document_id:int,user:stack.User=Depends(stack.me),s:Session=Depends(stack.db)):
    x=s.get(DocumentAsset,document_id)
    if not x or x.user_id!=user.id:raise HTTPException(404,"Document not found")
    path=(core.STORAGE/x.storage_key).resolve();root=core.STORAGE.resolve()
    if root in path.parents:path.unlink(missing_ok=True)
    s.delete(x);s.commit()
    return {"deleted":True,"id":document_id}


@app.post("/api/v1/documents/purge-expired")
def purge_expired(user:stack.User=Depends(stack.me),s:Session=Depends(stack.db)):
    rows=s.query(DocumentAsset).filter(DocumentAsset.user_id==user.id,DocumentAsset.retention_until<now()).all();deleted=[]
    for x in rows:
        path=(core.STORAGE/x.storage_key).resolve();root=core.STORAGE.resolve()
        if root in path.parents:path.unlink(missing_ok=True)
        deleted.append(x.id);s.delete(x)
    s.commit();return {"deleted_ids":deleted,"count":len(deleted)}
