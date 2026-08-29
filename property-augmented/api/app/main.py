from __future__ import annotations
import csv, hashlib, io, json, os, re, uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Literal
import httpx
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import Workbook, load_workbook
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

APP='Property Development, Augmented API'; VERSION='2.0.0'
OPENAI_API_KEY=os.getenv('OPENAI_API_KEY',''); OPENAI_MODEL=os.getenv('OPENAI_MODEL','gpt-4o-mini')
COMPANIES_HOUSE_API_KEY=os.getenv('COMPANIES_HOUSE_API_KEY',''); EPC_API_URL=os.getenv('EPC_API_URL',''); EPC_API_TOKEN=os.getenv('EPC_API_TOKEN','')
STORAGE=Path(os.getenv('STORAGE_DIR','./storage')); STORAGE.mkdir(parents=True,exist_ok=True)
CORS=[x.strip() for x in os.getenv('CORS_ORIGINS','http://localhost:3000').split(',') if x.strip()]
app=FastAPI(title=APP,version=VERSION,description='Independent data and AI backend. Lovable is not required.')
app.add_middleware(CORSMiddleware,allow_origins=CORS,allow_credentials=True,allow_methods=['*'],allow_headers=['*'])

def stamp(): return datetime.now(timezone.utc).isoformat()
def src(name,url,caveat=''): return {'name':name,'url':url,'retrieved_at':stamp(),'caveat':caveat}
def pc(v):
 s=re.sub(r'\s+','',v.upper()); return s[:-3]+' '+s[-3:] if len(s)>3 else v.upper().strip()
async def getjson(url,params=None,headers=None,auth=None,timeout=25):
 async with httpx.AsyncClient(timeout=timeout,follow_redirects=True) as c:
  r=await c.get(url,params=params,headers=headers,auth=auth); r.raise_for_status(); return r.json()

class SiteRequest(BaseModel):
 postcode:str; address:str=''; radius_km:float=Field(2,ge=.1,le=20); company_query:str=''; include_epc:bool=True
class Analysis(BaseModel):
 text:str=''; question:str=''; mode:str='planning-evidence'; context:dict[str,Any]=Field(default_factory=dict)
class Research(BaseModel):
 topic:str; postcode:str=''; address:str=''
class Appraisal(BaseModel):
 acquisition:float=0; transaction_costs:float=0; construction:float=0; professional_fees:float=0; finance:float=0; contingency:float=0; other_costs:float=0; gdv:float=0; rental_income_annual:float=0; holding_months:float=12
class Residual(BaseModel):
 gdv:float; construction:float; professional_fees:float=0; finance:float=0; contingency:float=0; other_costs:float=0; target_profit:float=0; transaction_costs:float=0
class Finance(BaseModel):
 principal:float; annual_rate_pct:float; months:float; compounding:Literal['simple','monthly']='simple'
class Quotes(BaseModel): quotes:list[dict[str,Any]]; context:str=''
class SEO(BaseModel): seeds:list[str]; audience:str='UK property developers and construction SMEs'; site_url:str=''
class Report(BaseModel): title:str; format:Literal['pdf','docx','xlsx','csv','json']='pdf'; data:dict[str,Any]

MODES={
'site-analyst':'Separate confirmed facts, assumptions, unknowns and material-if-wrong issues. Never decide whether to buy. Cite supplied sources; missing evidence = not established.',
'planning-evidence':'You are a UK planning evidence analyst, not a planning authority. Build evidence matrices, chronologies, contradictions, gaps and questions. Never predict approval.',
'procurement':'Normalise quotations, preserve bidder wording, flag exclusions/provisional sums/qualifications. Do not appoint or certify.',
'project-controls':'Structure risks, variations and decisions. Never infer approval; use approval not evidenced.',
'report-writer':'Distinguish fact, assumption, professional opinion and decision. Never invent evidence or credentials.',
'seo-strategist':'Create intent-led UK SEO clusters without stuffing. Never invent keyword volume or difficulty.'}
async def ai(mode,prompt,context=None):
 if not OPENAI_API_KEY:return {'configured':False,'content':'AI provider not configured; set OPENAI_API_KEY on the independent backend.'}
 payload={'model':OPENAI_MODEL,'messages':[{'role':'system','content':MODES.get(mode,MODES['report-writer'])},{'role':'user','content':json.dumps(context or {},default=str)[:90000]+'\n\nTASK\n'+prompt}],'temperature':.2}
 async with httpx.AsyncClient(timeout=90) as c:
  r=await c.post('https://api.openai.com/v1/chat/completions',headers={'Authorization':f'Bearer {OPENAI_API_KEY}','Content-Type':'application/json'},json=payload)
  if r.status_code>=400: raise HTTPException(502,f'AI provider error {r.status_code}')
  d=r.json(); return {'configured':True,'content':d['choices'][0]['message']['content'],'model':d.get('model'),'usage':d.get('usage')}

@app.get('/health')
def health(): return {'status':'healthy','service':APP,'version':VERSION,'ai_configured':bool(OPENAI_API_KEY)}
@app.get('/api/v1/data/sources')
def sources(): return {'sources':[src('Planning Data','https://www.planning.data.gov.uk/docs','Beta; coverage/completeness vary.'),src('HM Land Registry Price Paid Data','https://www.gov.uk/government/statistical-data-sets/price-paid-data-downloads','Registration lag/exclusions; not a valuation.'),src('Environment Agency Flood Monitoring','https://environment.data.gov.uk/flood-monitoring/doc/reference','No guaranteed SLA; not safety-critical advice.'),src('Energy Performance Data','https://get-energy-performance-data.communities.gov.uk/','Current service requires account/API setup.'),src('Companies House','https://developer.company-information.service.gov.uk/','API key required.'),src('GOV.UK','https://www.gov.uk/','Verify current policy wording.') ]}

async def geocode(postcode):
 u=f"https://api.postcodes.io/postcodes/{pc(postcode).replace(' ','%20')}"; d=await getjson(u); r=d.get('result')
 if not r: raise HTTPException(422,'Postcode could not be geocoded')
 return {'postcode':r.get('postcode'),'latitude':r.get('latitude'),'longitude':r.get('longitude'),'admin_district':r.get('admin_district'),'region':r.get('region'),'country':r.get('country'),'source':src('postcodes.io',u,'Postcode centroid, not surveyed site geometry.')}
async def planning(lat,lon):
 u='https://www.planning.data.gov.uk/entity.json'; params=[('latitude',lat),('longitude',lon),('limit',250)]
 for ds in ['planning-application','conservation-area','listed-building','green-belt','flood-risk-zone','article-4-direction-area','tree-preservation-zone','brownfield-site','ancient-woodland']: params.append(('dataset',ds))
 try:return {'data':await getjson(u,params=params),'source':src('Planning Data API',u,'Government Planning Data is beta; coverage varies by dataset/authority.')}
 except Exception as e:return {'data':None,'error':str(e),'source':src('Planning Data API',u,'Unavailable/incomplete for this request.')}
async def floods(lat,lon,dist):
 u='https://environment.data.gov.uk/flood-monitoring/id/floods'
 try:return {'data':await getjson(u,params={'lat':lat,'long':lon,'dist':dist}),'source':src('Environment Agency Flood Monitoring',u,'Real-time open data; not safety-critical advice.')}
 except Exception as e:return {'data':None,'error':str(e),'source':src('Environment Agency Flood Monitoring',u)}
async def hmlr(postcode):
 u='https://landregistry.data.gov.uk/landregistry/query'; p=pc(postcode); q=f'''PREFIX lrppi:<http://landregistry.data.gov.uk/def/ppi/> PREFIX lrcommon:<http://landregistry.data.gov.uk/def/common/> SELECT ?transaction ?price ?date ?propertyType ?paon ?street ?town ?postcode WHERE {{ ?transaction lrppi:pricePaid ?price;lrppi:transactionDate ?date;lrppi:propertyAddress ?address. OPTIONAL{{?transaction lrppi:propertyType ?propertyType.}} ?address lrcommon:postcode ?postcode. FILTER(UCASE(STR(?postcode))="{p}") OPTIONAL{{?address lrcommon:paon ?paon.}} OPTIONAL{{?address lrcommon:street ?street.}} OPTIONAL{{?address lrcommon:town ?town.}} }} ORDER BY DESC(?date) LIMIT 100'''
 try:
  async with httpx.AsyncClient(timeout=30) as c:r=await c.get(u,params={'query':q},headers={'Accept':'application/sparql-results+json'});r.raise_for_status();d=r.json()
  rows=[{k:v.get('value') for k,v in b.items()} for b in d.get('results',{}).get('bindings',[])]
  return {'transactions':rows,'source':src('HM Land Registry Price Paid Data',u,'Registration lag/exclusions; not a valuation. Contains HMLR data under OGL v3.0; verify current attribution before publication.')}
 except Exception as e:return {'transactions':[],'error':str(e),'source':src('HM Land Registry Price Paid Data',u)}
async def govuk(query):
 u='https://www.gov.uk/api/search.json'
 try:
  d=await getjson(u,params={'q':query,'count':30}); return {'results':[{'title':x.get('title'),'description':x.get('description'),'link':'https://www.gov.uk'+x.get('link',''),'public_timestamp':x.get('public_timestamp')} for x in d.get('results',[])],'source':src('GOV.UK Search API',u,'Discovery only; verify authoritative current policy text.')}
 except Exception as e:return {'results':[],'error':str(e),'source':src('GOV.UK Search API',u)}

@app.post('/api/v1/site/intelligence')
async def site(r:SiteRequest):
 g=await geocode(r.postcode); import asyncio; pl,fl,lr,po=await asyncio.gather(planning(g['latitude'],g['longitude']),floods(g['latitude'],g['longitude'],r.radius_km),hmlr(r.postcode),govuk(f"planning policy {g.get('admin_district','')} {r.address or r.postcode}"))
 return {'query':r.model_dump(),'geocode':g,'planning':pl,'floods':fl,'land_registry':lr,'policy_discovery':po,'generated_at':stamp(),'human_review_gate':'Verify boundaries, current policy, technical constraints and professional conclusions.'}
@app.post('/api/v1/planning/analyse')
async def planning_analyse(r:Analysis): return {'analysis':await ai('planning-evidence',r.question or 'Build an evidence matrix, chronology, contradictions, missing evidence and professional questions. Cite supplied sources.',{'material':r.text,'context':r.context})}
@app.post('/api/v1/planning/refusal')
async def refusal(r:Analysis): return {'analysis':await ai('planning-evidence','Break each refusal reason into policy, factual premise, technical premise, judgement, evidence needed, professional discipline and clarification question. Do not predict appeal success.',{'material':r.text})}
@app.post('/api/v1/ai/assist')
async def assist(r:Analysis): return await ai(r.mode,r.question or r.text,r.context)
@app.post('/api/v1/research/deep')
async def research(r:Research):
 bundle={'topic':r.topic,'govuk':await govuk(r.topic)}
 if r.postcode:bundle['site_intelligence']=await site(SiteRequest(postcode=r.postcode,address=r.address))
 return {'status':'completed','analysis':await ai('report-writer','Produce a source-backed research brief. Separate established facts, interpretation, uncertainty, missing evidence and verification. Do not fabricate sources.',bundle),'source_bundle':bundle,'completed_at':stamp()}

@app.post('/api/v1/calculators/appraisal')
def appraisal(r:Appraisal):
 total=r.acquisition+r.transaction_costs+r.construction+r.professional_fees+r.finance+r.contingency+r.other_costs; profit=r.gdv-total
 return {'total_development_cost':total,'profit':profit,'margin_on_cost_pct':profit/total*100 if total else None,'margin_on_gdv_pct':profit/r.gdv*100 if r.gdv else None,'roi_pct':profit/total*100 if total else None,'gross_yield_on_gdv_pct':r.rental_income_annual/r.gdv*100 if r.gdv and r.rental_income_annual else None,'assumptions':r.model_dump(),'note':'Arithmetic decision support, not a valuation, tax calculation, QS cost plan or investment recommendation.'}
@app.post('/api/v1/calculators/residual-land-value')
def residual(r:Residual):
 deduct=r.construction+r.professional_fees+r.finance+r.contingency+r.other_costs+r.target_profit+r.transaction_costs;return {'residual_land_value':r.gdv-deduct,'gdv':r.gdv,'deductions':deduct,'note':'Residual arithmetic only; verify inputs.'}
@app.post('/api/v1/calculators/finance')
def finance(r:Finance):
 rate=r.annual_rate_pct/100; amount=r.principal*((1+rate/12)**r.months) if r.compounding=='monthly' else r.principal*(1+rate*r.months/12);return {'interest':amount-r.principal,'total':amount,'method':r.compounding,'note':'Illustrative arithmetic; excludes lender-specific fees/drawdowns.'}
@app.post('/api/v1/procurement/compare-quotes')
async def quotes(r:Quotes): return {'analysis':await ai('procurement','Normalise these quotes into comparable packages; identify inclusions, exclusions, quantities, provisional sums, VAT/programme/qualifications and clarification questions.',r.model_dump())}

BASE=['AI for property developers UK','AI construction software UK','property development feasibility software','planning application AI','planning policy analysis','property development appraisal calculator','construction quote comparison','construction risk register','variation tracker construction','development due diligence checklist','planning evidence matrix','AI project management construction','UK property data analysis','small housebuilder AI','planning constraints checker','property development consultant AI','construction workflow automation']
@app.post('/api/v1/seo/keywords')
async def seo(r:SEO):
 seeds=list(dict.fromkeys(r.seeds+BASE)); return {'seeds':seeds,'analysis':await ai('seo-strategist','Build a UK intent-led keyword/content map: target route, intent, page type, H1, title, meta, internal links and content gaps. Search volume/difficulty must be null unless measured data is supplied.',{'seeds':seeds,'audience':r.audience,'site':r.site_url}),'note':'No ranking guarantee and no invented keyword metrics.'}

@app.post('/api/v1/documents/upload')
async def upload(file:UploadFile=File(...),project_id:str=Form(default='')):
 raw=await file.read();
 if len(raw)>30*1024*1024:raise HTTPException(413,'30MB limit')
 name=re.sub(r'[^A-Za-z0-9._-]+','_',file.filename or 'upload'); path=STORAGE/f'{uuid.uuid4().hex}_{name}';path.write_bytes(raw);suffix=path.suffix.lower();text=''
 try:
  if suffix=='.pdf':text='\n'.join((p.extract_text() or '') for p in PdfReader(str(path)).pages)
  elif suffix=='.docx':text='\n'.join(p.text for p in DocxDocument(str(path)).paragraphs)
  elif suffix in {'.xlsx','.xlsm'}:
   wb=load_workbook(path,data_only=True,read_only=True);text='\n'.join('\t'.join('' if v is None else str(v) for v in row) for ws in wb.worksheets for row in ws.iter_rows(values_only=True))
  elif suffix in {'.csv','.txt','.md','.json'}:text=path.read_text(encoding='utf-8',errors='replace')
 except Exception as e:text=f'[Extraction error: {e}]'
 return {'id':path.name,'filename':file.filename,'sha256':hashlib.sha256(raw).hexdigest(),'text_chars':len(text),'text':text[:150000],'project_id':project_id,'source':{'type':'user-upload','filename':file.filename,'retrieved_at':stamp()}}

@app.post('/api/v1/reports/generate')
def report(r:Report):
 if r.format=='json':raw=json.dumps(r.data,indent=2,default=str).encode();mime='application/json';fn='report.json'
 elif r.format=='csv':
  s=io.StringIO();w=csv.writer(s);w.writerow(['Field','Value']);[w.writerow([k,json.dumps(v,default=str) if isinstance(v,(dict,list)) else v]) for k,v in r.data.items()];raw=s.getvalue().encode();mime='text/csv';fn='report.csv'
 elif r.format=='xlsx':
  wb=Workbook();ws=wb.active;ws.append([r.title]);ws.append(['Field','Value']);[ws.append([k,json.dumps(v,default=str) if isinstance(v,(dict,list)) else v]) for k,v in r.data.items()];b=io.BytesIO();wb.save(b);raw=b.getvalue();mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet';fn='report.xlsx'
 elif r.format=='docx':
  d=DocxDocument();d.add_heading(r.title,0);[d.add_paragraph(f'{k}: {json.dumps(v,default=str) if isinstance(v,(dict,list)) else v}') for k,v in r.data.items()];b=io.BytesIO();d.save(b);raw=b.getvalue();mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document';fn='report.docx'
 else:
  b=io.BytesIO();styles=getSampleStyleSheet();story=[Paragraph(r.title,styles['Title']),Paragraph('Property Development, Augmented | SITE → BUILD → PROVE',styles['Heading2']),Spacer(1,12)];[story.extend([Paragraph(str(k).replace('_',' ').title(),styles['Heading2']),Paragraph(str(v)[:5000],styles['BodyText'])]) for k,v in r.data.items()];SimpleDocTemplate(b,pagesize=A4).build(story);raw=b.getvalue();mime='application/pdf';fn='report.pdf'
 return StreamingResponse(io.BytesIO(raw),media_type=mime,headers={'Content-Disposition':f'attachment; filename="{fn}"'})
